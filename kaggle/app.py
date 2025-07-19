import os
import sys
import logging
import hashlib
import json
import base64
import subprocess
import shutil
import tempfile
import fcntl
import zipfile
import mimetypes
from datetime import datetime, timezone
from functools import wraps
from pathlib import Path
from collections import Counter
from flask import (
    Flask,
    request,
    jsonify,
    render_template,
    send_file,
    redirect,
    url_for,
    current_app,
)
from werkzeug.utils import secure_filename
from celery import Celery
from docx import Document as DocxWriter
from docx import Document as DocxDoc
from pptx import Presentation
from fpdf import FPDF
import pytesseract
from PIL import Image
import pdfplumber
import fitz
import pandas as pd
import numpy as np
import faiss
from openai import OpenAI
import replicate
from replicate.exceptions import ReplicateError
from filelock import FileLock
from rag_pipeline import generate_media

INDEX_VERSION_FILE = "index_version.json"  # For version tracking

print("Starting Crossbar Flask app...")


def make_celery(app):
    celery = Celery(
        app.import_name,
        backend=app.config["CELERY_RESULT_BACKEND"],
        broker=app.config["CELERY_BROKER_URL"],
    )
    celery.conf.update(app.config)
    return celery


app = Flask(__name__)
app.config.update(
    CELERY_BROKER_URL="redis://localhost:6379/0",
    CELERY_RESULT_BACKEND="redis://localhost:6379/0",
)
celery = make_celery(app)


def _to_native(o):
    # teach json how to handle NumPy numbers
    if isinstance(o, (np.float_, np.floating)):
        return float(o)
    if isinstance(o, (np.int_, np.integer)):
        return int(o)
    raise TypeError  # let real errors propagate


VALID_RATIOS = ["16:9", "9:16", "4:3", "1:1", "3:2", "2:3"]
VALID_STYLES = ["AUTO", "GENERAL", "REALISTIC", "DESIGN"]

# Configuration
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
ADMIN_OUTPUTS = "admin_outputs"
CONVERSATIONS_FILE = "conversations.json"
CONTEXT_UPLOADS_FILE = "context_uploads.json"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(ADMIN_OUTPUTS, exist_ok=True)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("DocumentAIApp")

# Supported file types
ALLOWED_EXTENSIONS = {
    ".txt",
    ".pdf",
    ".doc",
    ".docx",
    ".ppt",
    ".pptx",
    ".xls",
    ".xlsx",
    ".csv",
    ".tsv",
    ".jpg",
    ".jpeg",
    ".png",
    ".odt",
    ".rtf",
    ".zip",
}

# Admin credentials
ADMIN_CREDENTIALS = {
    "username": os.environ.get("ADMIN_USER", "admin"),
    "password": os.environ.get("ADMIN_PASSWORD", "admin"),
}

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100MB limit


# FAISS Vector Index with file locking
class PersistentVectorIndex:
    def __init__(self, path="vector_index"):
        self.path = path
        self.index_path = f"{path}.faiss"
        self.meta_path = f"{path}.json"
        self.lock_path = f"{path}.lock"

        # Initialize version tracking
        self._init_index_version()

        if os.path.exists(self.index_path):
            with self._acquire_lock():
                self.index = faiss.read_index(self.index_path)
                with open(self.meta_path, "r") as f:
                    self.metadata = json.load(f)
        else:
            self.index = faiss.IndexIDMap(faiss.IndexFlatL2(3072))
            self.metadata = {}

    def _init_index_version(self):
        """Initialize index version tracking"""
        if not os.path.exists(INDEX_VERSION_FILE):
            with open(INDEX_VERSION_FILE, "w") as f:
                json.dump({"version": 1}, f)

    def _acquire_lock(self):
        """Acquire a file lock for thread-safe operations"""
        lock = open(self.lock_path, "w")
        fcntl.flock(lock, fcntl.LOCK_EX)
        return lock

    def _release_lock(self, lock):
        """Release the file lock"""
        fcntl.flock(lock, fcntl.LOCK_UN)
        lock.close()

    def _increment_version(self):
        """Increment the index version"""
        with self._acquire_lock():
            try:
                with open(INDEX_VERSION_FILE, "r+") as f:
                    version_data = json.load(f)
                    version_data["version"] += 1
                    f.seek(0)
                    json.dump(version_data, f)
                    f.truncate()
                return version_data["version"]
            except Exception as e:
                logger.error(f"Error incrementing index version: {str(e)}")
                return None

    def get_current_version(self):
        """Get current index version"""
        try:
            with open(INDEX_VERSION_FILE, "r") as f:
                return json.load(f).get("version", 1)
        except:
            return 1

    def save(self):
        """Save index with version check"""
        current_version = self.get_current_version()
        with self._acquire_lock():
            try:
                # Check if index needs rebuild
                if not os.path.exists(self.index_path) or self._needs_rebuild():
                    faiss.write_index(self.index, self.index_path)
                    with open(self.meta_path, "w") as f:
                        json.dump(self.metadata, f, indent=2)
                    self._increment_version()
            except Exception as e:
                logger.error(f"Error saving index: {str(e)}")
                raise

    def _needs_rebuild(self):
        """Check if index needs rebuild based on metadata changes"""
        if not os.path.exists(self.meta_path):
            return True

        try:
            with open(self.meta_path, "r") as f:
                saved_metadata = json.load(f)
                return saved_metadata != self.metadata
        except:
            return True

    def add_vector(self, vector, metadata):
        """Add vector with thread-safe operation"""
        with self._acquire_lock():
            vector_id = len(self.metadata) + 1
            self.index.add_with_ids(
                np.array([vector]).astype("float32"),
                np.array([vector_id], dtype=np.int64),
            )
            self.metadata[vector_id] = metadata
            return vector_id

    def remove_vector(self, vector_id):
        """Remove vector with thread-safe operation"""
        with self._acquire_lock():
            self.index.remove_ids(np.array([vector_id], dtype=np.int64))
            if vector_id in self.metadata:
                del self.metadata[vector_id]

    def search(self, vector, k=5):
        """Search index with thread-safe operation"""
        with self._acquire_lock():
            distances, vector_ids = self.index.search(
                np.array([vector]).astype("float32"), k
            )
            results = []
            for i in range(len(vector_ids[0])):
                if vector_ids[0][i] >= 0 and vector_ids[0][i] in self.metadata:
                    results.append(
                        {
                            **self.metadata[vector_ids[0][i]],
                            "similarity": 1 - distances[0][i],
                        }
                    )
            return results


# Initialize the persistent index
vector_index = PersistentVectorIndex()


# Load existing context
def load_context_index():
    global vector_index
    if os.path.exists(CONTEXT_UPLOADS_FILE):
        try:
            with open(CONTEXT_UPLOADS_FILE, "r") as f:
                context_data = json.load(f)
                for entry in context_data:
                    if "embedding" in entry:
                        vector_index.add_vector(
                            np.array(entry["embedding"]),
                            {
                                "context_id": entry["context_id"],
                                "session_id": entry["session_id"],
                                "filename": entry.get("filename"),
                                "text": entry.get("text"),
                                "concepts": entry.get("concepts", []),
                                "priority": entry.get("priority", 0),
                                "timestamp": entry["timestamp"],
                            },
                        )
            # Save the loaded index to disk
            vector_index.save()
        except Exception as e:
            logger.error(f"Error loading context index: {str(e)}")


# Save context to JSON
def save_context_upload(entry):
    data = []
    if os.path.exists(CONTEXT_UPLOADS_FILE):
        with open(CONTEXT_UPLOADS_FILE, "r") as f:
            try:
                data = json.load(f)
            except:
                pass

    entry_copy = entry.copy()
    if "embedding" in entry_copy:
        del entry_copy["embedding"]
    data.append(entry_copy)

    with open(CONTEXT_UPLOADS_FILE, "w") as f:
        json.dump(data, f, indent=2)


# Helper function for tagging
def update_tags(filename, tags):
    """Update tags for a generated output file"""
    tags_file = os.path.join(ADMIN_OUTPUTS, "tags.json")
    tag_data = {}

    if os.path.exists(tags_file):
        try:
            with open(tags_file, "r") as f:
                tag_data = json.load(f)
        except:
            pass

    tag_data[filename] = tags

    with open(tags_file, "w") as f:
        json.dump(tag_data, f, indent=2)


# Utility functions
def allowed_file(filename):
    """Check if the file extension is allowed"""
    ext = Path(filename).suffix.lower()
    return ext in ALLOWED_EXTENSIONS


def hash_text(text):
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def is_image_heavy_pdf(filepath):
    try:
        doc = fitz.open(filepath)
        image_page_count = sum(1 for p in doc if len(p.get_images(full=True)) > 0)
        return image_page_count / len(doc) > 0.5
    except:
        return False


def extract_text_from_zip(zip_path):
    """Extract text from zip archive"""
    extracted_text = ""
    with zipfile.ZipFile(zip_path, "r") as zip_ref:
        for file in zip_ref.namelist():
            if any(file.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS):
                with zip_ref.open(file) as f:
                    try:
                        # Save to temp file for processing
                        with tempfile.NamedTemporaryFile(delete=False) as tmp:
                            tmp.write(f.read())
                            tmp_path = tmp.name

                        extracted_text += extract_text_from_file(tmp_path) + "\n\n"
                        os.unlink(tmp_path)
                    except Exception as e:
                        logger.error(f"Error processing {file} in ZIP: {str(e)}")
    return extracted_text


def extract_text_from_file(filepath, api_key=None):
    """Extract text from various file types using GPT-4o vision when needed"""
    ext = Path(filepath).suffix.lower()
    text = ""

    try:
        if ext == ".pdf":
            if is_image_heavy_pdf(filepath):
                # Use GPT-4o vision for image-heavy PDFs
                text = process_pdf_with_vision(filepath, api_key)
            else:
                # Use regular text extraction for text-based PDFs
                text = extract_pdf_text(filepath)

        elif ext in [".jpg", ".jpeg", ".png"]:
            # Use GPT-4o vision for images
            text = analyze_image_with_gpt4o(filepath, api_key)

        elif ext in [".doc", ".docx"]:
            text = extract_docx_text(filepath)

        elif ext == ".txt":
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()

    except Exception as e:
        logger.error(f"Error extracting text: {str(e)}")
        return ""

    return text


def process_pdf_with_vision(pdf_path, api_key):
    """Use GPT-4o vision to extract text from PDF pages"""
    client = OpenAI(api_key=api_key)
    doc = fitz.open(pdf_path)
    full_text = []

    for page_num in range(len(doc)):
        # Render page as image
        pix = doc[page_num].get_pixmap()
        img_path = f"page_{page_num}.png"
        pix.save(img_path)

        # Send to GPT-4o vision
        with open(img_path, "rb") as img_file:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": "Extract all text from this document page, preserving layout, tables and structure.",
                    },
                    {"role": "user", "image": img_file.read()},
                ],
                max_tokens=2000,
            )
            full_text.append(response.choices[0].message.content)

        # Clean up
        os.remove(img_path)

    return "\n\n".join(full_text)


def analyze_image_with_gpt4o(image_path, api_key):
    """Use GPT-4o vision to analyze images/charts"""
    client = OpenAI(api_key=api_key)
    with open(image_path, "rb") as img_file:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": "Describe this image in detail, including any text, charts, or diagrams.",
                },
                {"role": "user", "image": img_file.read()},
            ],
            max_tokens=1000,
        )
        return response.choices[0].message.content


def gpt_semantic_chunk(text, api_key, max_chunks=10):
    """Use GPT-4o to split text into coherent semantic chunks"""
    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": f"""
                    Split this document into up to {max_chunks} coherent semantic chunks.
                    Each chunk should represent a logically self-contained section or idea (~400-600 words).
                    Preserve all formatting, tables, and layout information.
                    Return a JSON list of chunks.
                    
                    Document:
                    {text[:6000]}  # Limit context window
                    """,
                }
            ],
            response_format={"type": "json_object"},
            temperature=0.3,
        )
        return json.loads(response.choices[0].message.content)["chunks"]
    except Exception as e:
        logger.error(f"GPT chunking failed: {str(e)}")
        return [text]  # Fallback to original text


def update_context_priority(session_id, used_context_ids, positive=True):
    """Adjust priority based on usage"""
    if not os.path.exists(CONTEXT_UPLOADS_FILE):
        return

    try:
        with open(CONTEXT_UPLOADS_FILE, "r+") as f:
            data = json.load(f)
            for entry in data:
                if entry["context_id"] in used_context_ids:
                    # Increase priority for positive feedback, decrease for negative
                    entry["priority"] += 1 if positive else -1
                    # Ensure priority stays within reasonable bounds
                    entry["priority"] = max(0, min(entry["priority"], 10))

            f.seek(0)
            json.dump(data, f, indent=2)
            f.truncate()

    except Exception as e:
        logger.error(f"Error updating context priorities: {str(e)}")


def embed(text, api_key, model="text-embedding-3-large"):
    client = OpenAI(api_key=api_key)
    return (
        client.embeddings.create(model=model, input=[text], encoding_format="float")
        .data[0]
        .embedding
    )


def get_enhanced_embeddings(texts, api_key, model="text-embedding-3-large"):
    """Generate embeddings with cross-chunk context"""
    if not isinstance(texts, list):
        texts = [texts]

    try:
        client = OpenAI(api_key=api_key)

        # First get a global summary of all chunks
        summary_prompt = f"""
        Analyze these text chunks as a cohesive document and create a global summary 
        that captures the main themes and relationships between chunks:
        {texts}
        """

        summary = (
            client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": summary_prompt}],
                max_tokens=500,
            )
            .choices[0]
            .message.content
        )

        # Get embeddings for each chunk + summary
        embeddings = []
        for text in texts:
            # Combine chunk text with global summary
            enhanced_text = f"Chunk: {text}\n\nGlobal Context: {summary}"
            embedding = embed(enhanced_text, api_key, model)
            embeddings.append(embedding)

        return embeddings[0] if len(embeddings) == 1 else embeddings

    except Exception as e:
        logger.error(f"Enhanced embeddings failed: {str(e)}")
        return embed(texts, api_key, model)  # Fallback to regular embedding


def detect_document_type(text, api_key):
    """Use GPT to classify document type"""
    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": f"""
                    Classify this document into one of these types:
                    [memo, article, report, slide, code, legal, spreadsheet, general]
                    
                    Text:
                    {text[:2000]}
                    
                    Return only the type.
                    """,
                }
            ],
            max_tokens=10,
            temperature=0,
        )
        return response.choices[0].message.content.strip().lower()
    except Exception as e:
        logger.error(f"Document type detection failed: {str(e)}")
        return "general"


def extract_concepts(text, api_key):
    client = OpenAI(api_key=api_key)
    prompt = (
        "Extract key concepts and summarize the main ideas from the following text:\n"
        + text[:4000]
    )
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=300,
    )
    return response.choices[0].message.content


def save_conversation(session_id, question, answer):
    history = []
    if os.path.exists(CONVERSATIONS_FILE):
        with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
            history = json.load(f)

    history.append(
        {
            "session_id": session_id,
            "question": question,
            "answer": answer,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }
    )

    with open(CONVERSATIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2)


def load_conversation():
    if os.path.exists(CONVERSATIONS_FILE):
        with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def save_context_upload_with_priority(
    session_id, filename=None, text=None, concepts=[], priority=0
):
    context_id = hashlib.md5(
        f"{session_id}{datetime.now(timezone.utc).isoformat()}".encode()
    ).hexdigest()

    entry = {
        "context_id": context_id,
        "session_id": session_id,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "concepts": concepts,
        "priority": priority,
    }

    if filename:
        entry["filename"] = filename
        entry["type"] = "file"
    elif text:
        entry["text"] = text[:100] + "..." if len(text) > 100 else text
        entry["type"] = "text"

    save_context_upload(entry)
    return context_id


def trigger_index_rebuild():
    """Delete existing index to trigger rebuild on next question"""
    index_path = os.path.join(OUTPUT_FOLDER, "index.faiss")
    metadata_path = os.path.join(OUTPUT_FOLDER, "metadata.json")

    try:
        if os.path.exists(index_path):
            os.remove(index_path)
        if os.path.exists(metadata_path):
            os.remove(metadata_path)
        logger.info("Index files removed, will rebuild on next question")
    except Exception as e:
        logger.error(f"Error removing index files: {str(e)}")


# Authentication decorator
def requires_admin_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        if not auth or not (
            auth.username == ADMIN_CREDENTIALS["username"]
            and auth.password == ADMIN_CREDENTIALS["password"]
        ):
            return jsonify({"error": "Unauthorized"}), 401
        return f(*args, **kwargs)

    return decorated


# API Routes
@app.route("/api/admin_login", methods=["POST"])
def admin_login():
    data = request.json
    if (
        data.get("username") == ADMIN_CREDENTIALS["username"]
        and data.get("password") == ADMIN_CREDENTIALS["password"]
    ):
        return jsonify({"status": "success"})
    return jsonify({"error": "Unauthorized"}), 403


@app.route("/api/admin/reset_db", methods=["POST"])
@requires_admin_auth
def reset_db():
    try:
        shutil.rmtree(UPLOAD_FOLDER, ignore_errors=True)
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        shutil.rmtree(OUTPUT_FOLDER, ignore_errors=True)
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)

        if os.path.exists(CONVERSATIONS_FILE):
            os.remove(CONVERSATIONS_FILE)
        if os.path.exists(CONTEXT_UPLOADS_FILE):
            os.remove(CONTEXT_UPLOADS_FILE)

        # Reset vector index
        vector_index.index.reset()
        vector_index.metadata = {}
        vector_index.save()  # Save the empty index

        return jsonify({"status": "reset complete"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ask", methods=["POST"])
def api_ask():
    data = request.json
    question = data.get("question")
    session_id = data.get("sessionId")
    model = data.get("model", "text-embedding-3-large")
    role = data.get("sessionProfile", "general")
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")
    enhance = data.get("enhance", False)

    if not question or not api_key:
        return jsonify({"error": "Missing question or API key"}), 400

    try:
        # 1. First try to answer from local context
        q_emb = embed(question, api_key, model)
        local_results = vector_index.search(q_emb, k=5)

        # Build context from local files
        local_context = ""
        matched_chunks = []
        for i, res in enumerate(local_results[:3]):  # Use top 3 local results
            if res.get("text"):
                local_context += f"[Local Source {i+1}]: {res['text']}\n\n"
                matched_chunks.append(
                    {
                        "file": res.get("filename", "text snippet"),
                        "text": (
                            res["text"][:500] + "..."
                            if len(res["text"]) > 500
                            else res["text"]
                        ),
                        "similarity": float(res["similarity"]),
                        "concepts": res.get("concepts", []),
                    }
                )

        # 2. Always perform web search (simulated or real)
        try:
            web_results = perform_web_search(question, api_key)
            web_context = "\n".join(
                f"[Web Source {i+1}]: {res['snippet']} (from {res['source']})"
                for i, res in enumerate(web_results[:2])  # Use top 2 web results
            )
        except Exception as e:
            logger.warning(f"Web search failed: {str(e)}")
            web_context = ""

        # 3. Combine contexts with priority to local files
        combined_context = ""
        if local_context:
            combined_context += (
                f"LOCAL CONTEXT (from uploaded files):\n{local_context}\n"
            )
        if web_context:
            combined_context += f"WEB CONTEXT (from internet search):\n{web_context}\n"

        # 4. Generate answer
        client = OpenAI(api_key=api_key)
        prompt = f"""
        [Role: {role}]
        You have access to both local context (from user-uploaded files) and web context.
        Prioritize the local context when it contains relevant information, as it may be 
        more specific or up-to-date than web sources. Only use web context when:
        - The local context doesn't contain the answer
        - The web information is clearly more current (e.g., for news, recent events)
        - You need to supplement local information with general knowledge
        
        {combined_context}
        
        Question: {question}
        Answer:
        """

        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1024,
            temperature=0.3,
        )
        answer = response.choices[0].message.content.strip()

        # 5. Format response
        response_data = {
            "answer": answer,
            "question": question,
            "matched_chunks": matched_chunks,
            "model": model,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }

        return jsonify(response_data)

    except Exception as e:
        logger.error(f"Error in api_ask: {str(e)}", exc_info=True)
        return (
            jsonify({"error": "An error occurred while processing your request"}),
            500,
        )


def perform_web_search(query, api_key):
    """Simulate or perform actual web search"""
    client = OpenAI(api_key=api_key)

    # For a real implementation, you would use a search API here
    # This is a simulation that generates realistic-looking web results

    prompt = f"""
    You are a web search assistant. Given the query "{query}", generate 3-5 realistic 
    search results with titles, URLs, and snippets as if they came from a real search engine.
    
    Return the results in JSON format like this:
    {{
        "results": [
            {{
                "title": "Result title",
                "source": "example.com",
                "snippet": "Relevant information snippet..."
            }}
        ]
    }}
    """

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
    )

    data = json.loads(response.choices[0].message.content)
    return data.get("results", [])


@app.route("/upload", methods=["POST"])
def upload_file():
    if "files" not in request.files:
        return jsonify({"error": "No files provided"}), 400

    files = request.files.getlist("files")
    session_id = request.form.get("sessionId")
    api_key = request.form.get("apiKey")
    model = request.form.get("model", "text-embedding-3-large")

    if not api_key:
        return jsonify({"error": "API key required"}), 400

    results = []
    seen_hashes = set()

    for file in files:
        if file.filename == "":
            continue

        if not allowed_file(file.filename):
            continue

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(filepath)

        try:
            # Extract text
            text = extract_text_from_file(filepath)
            if not text.strip():
                continue

            h = hash_text(text)
            if h in seen_hashes:
                continue
            seen_hashes.add(h)

            # Embed and extract concepts
            embedding = get_enhanced_embeddings(text, api_key, model)
            concepts_raw = extract_concepts(text, api_key)
            # Ensure concepts is always a list
            if isinstance(concepts_raw, str):
                concepts = [c.strip() for c in concepts_raw.split(",") if c.strip()]
            else:
                concepts = []

            # Add to vector index
            metadata = {
                "context_id": hashlib.md5(
                    f"{session_id}{datetime.now(timezone.utc).isoformat()}".encode()
                ).hexdigest(),
                "session_id": session_id,
                "filename": filename,
                "text": text,
                "concepts": concepts,
                "priority": 1,  # File context has higher priority
                "timestamp": datetime.now(timezone.utc).isoformat(),
            }
            vector_id = vector_index.add_vector(embedding, metadata)

            # Save to context uploads
            entry = {
                **metadata,
                "embedding": embedding,  # Removed .tolist()
            }
            save_context_upload(entry)

            results.append(
                {
                    "filename": filename,
                    "status": "processed",
                    "concepts": concepts,
                }
            )

        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}")
            results.append(
                {
                    "filename": filename,
                    "status": "error",
                    "error": str(e),
                }
            )
            continue

    trigger_index_rebuild()
    return jsonify({"status": "success", "results": results})


@app.route("/api/upload_context", methods=["POST"])
def upload_context():
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not api_key:
        return jsonify({"error": "Missing API key"}), 401

    model = request.form.get("model", "text-embedding-3-large")
    session_id = request.form.get("sessionId")
    if not session_id:
        return jsonify({"error": "Missing sessionId"}), 400

    results = []
    seen_hashes = set()

    if "text" in request.form:
        text = request.form["text"]
        h = hash_text(text)
        if h not in seen_hashes:
            seen_hashes.add(h)
            try:
                embedding = get_enhanced_embeddings(text, api_key, model)
                concepts_raw = extract_concepts(text, api_key)
                # Ensure concepts is always a list
                if isinstance(concepts_raw, str):
                    concepts = [c.strip() for c in concepts_raw.split(",") if c.strip()]
                else:
                    concepts = []

                # Add to vector index
                metadata = {
                    "context_id": hashlib.md5(
                        f"{session_id}{datetime.now(timezone.utc).isoformat()}".encode()
                    ).hexdigest(),
                    "session_id": session_id,
                    "filename": None,
                    "text": text,
                    "concepts": concepts,
                    "priority": 0,  # Text context has lower priority
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                }
                vector_id = vector_index.add_vector(embedding, metadata)

                # Save to context uploads
                entry = {
                    **metadata,
                    "embedding": embedding,  # Removed .tolist()
                }
                save_context_upload(entry)

                results.append(
                    {
                        "context_id": metadata["context_id"],
                        "text": text,
                        "concepts": concepts,
                    }
                )

            except Exception as e:
                return jsonify({"error": str(e)}), 500

    elif "filesBase64" in request.form:
        files = request.form.getlist("filesBase64")
        filenames = request.form.getlist("filenames")

        for i, b64 in enumerate(files):
            try:
                _, encoded = b64.split(",", 1)
                file_bytes = base64.b64decode(encoded)
                filename = secure_filename(filenames[i])
                path = os.path.join(UPLOAD_FOLDER, filename)

                with open(path, "wb") as f:
                    f.write(file_bytes)

                # Extract text
                text = extract_text_from_file(path)
                h = hash_text(text)
                if h in seen_hashes:
                    continue
                seen_hashes.add(h)

                # Embed and extract concepts
                embedding = get_enhanced_embeddings(text, api_key, model)
                concepts_raw = extract_concepts(text, api_key)
                # Ensure concepts is always a list
                if isinstance(concepts_raw, str):
                    concepts = [c.strip() for c in concepts_raw.split(",") if c.strip()]
                else:
                    concepts = []

                # Add to vector index
                metadata = {
                    "context_id": hashlib.md5(
                        f"{session_id}{datetime.now(timezone.utc).isoformat()}".encode()
                    ).hexdigest(),
                    "session_id": session_id,
                    "filename": filename,
                    "text": text,
                    "concepts": concepts,
                    "priority": 1,  # File context has higher priority
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                }
                vector_id = vector_index.add_vector(embedding, metadata)

                # Save to context uploads
                entry = {
                    **metadata,
                    "embedding": embedding,  # Removed .tolist()
                }
                save_context_upload(entry)

                results.append(
                    {
                        "context_id": metadata["context_id"],
                        "filename": filename,
                        "concepts": concepts,
                    }
                )

            except Exception as e:
                logger.error(f"Error processing file {filenames[i]}: {str(e)}")
                continue

    return jsonify({"status": "success", "results": results})


@app.route("/api/list_contexts", methods=["GET"])
def list_contexts():
    session_id = request.args.get("sessionId")
    if not session_id:
        return jsonify({"error": "Missing sessionId"}), 400

    contexts = []
    if os.path.exists(CONTEXT_UPLOADS_FILE):
        try:
            with open(CONTEXT_UPLOADS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                contexts = [
                    entry for entry in data if entry.get("session_id") == session_id
                ]
        except Exception as e:
            logger.error(f"Error loading context uploads: {str(e)}")

    return jsonify({"status": "success", "contexts": contexts})


@app.route("/api/remove_context", methods=["POST"])
def remove_context():
    data = request.json
    context_id = data.get("context_id")
    session_id = data.get("session_id")
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")

    if not context_id or not session_id:
        return jsonify({"error": "Missing context_id or session_id"}), 400

    # Remove from vector index
    vector_id_to_remove = None
    for vec_id, meta in vector_index.metadata.items():
        if meta["context_id"] == context_id and meta["session_id"] == session_id:
            vector_id_to_remove = vec_id
            break

    if vector_id_to_remove:
        vector_index.remove_vector(vector_id_to_remove)

    # Update context_uploads.json
    if os.path.exists(CONTEXT_UPLOADS_FILE):
        try:
            with open(CONTEXT_UPLOADS_FILE, "r", encoding="utf-8") as f:
                context_data = json.load(f)
            context_data = [e for e in context_data if e["context_id"] != context_id]
            with open(CONTEXT_UPLOADS_FILE, "w", encoding="utf-8") as f:
                json.dump(context_data, f, indent=2)
        except Exception as e:
            logger.error(f"Error removing context from file: {e}")

    return jsonify({"status": "success"})


@app.route("/api/download_conversation", methods=["GET"])
def download():
    fmt = request.args.get("format", "txt")
    session_id = request.args.get("sessionId")

    if not session_id:
        return jsonify({"error": "Missing sessionId"}), 400

    convo = [c for c in load_conversation() if c["session_id"] == session_id]
    if not convo:
        return jsonify({"error": "No data"}), 404

    content = "\n\n".join([f"Q: {c['question']}\nA: {c['answer']}" for c in convo])

    if fmt == "txt":
        return content, 200, {"Content-Type": "text/plain"}
    elif fmt == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in content.splitlines():
            pdf.cell(200, 10, txt=line, ln=1)
        pdf_path = "conversation.pdf"
        pdf.output(pdf_path)
        return send_file(pdf_path, as_attachment=True)
    elif fmt == "docx":
        doc = DocxWriter()
        for q in convo:
            doc.add_paragraph(f"Q: {q['question']}")
            doc.add_paragraph(f"A: {q['answer']}\n")
        doc_path = "conversation.docx"
        doc.save(doc_path)
        return send_file(doc_path, as_attachment=True)
    else:
        return jsonify({"error": "invalid format"}), 400


# Admin routes
@app.route("/api/admin/uploads", methods=["GET"])
@requires_admin_auth
def admin_uploads():
    if os.path.exists(CONTEXT_UPLOADS_FILE):
        with open(CONTEXT_UPLOADS_FILE, "r") as f:
            return jsonify(json.load(f))
    return jsonify([])


@app.route("/api/admin/conversations", methods=["GET"])
@requires_admin_auth
def admin_conversations():
    if os.path.exists(CONVERSATIONS_FILE):
        with open(CONVERSATIONS_FILE, "r") as f:
            return jsonify(json.load(f))
    return jsonify([])


@app.route("/api/admin/frequent_questions", methods=["GET"])
@requires_admin_auth
def frequent_questions():
    n = int(request.args.get("n", 10))
    if os.path.exists(CONVERSATIONS_FILE):
        with open(CONVERSATIONS_FILE, "r") as f:
            conversations = json.load(f)
            questions = [c["question"] for c in conversations]
            counter = Counter(questions)
            top_n = counter.most_common(n)
            return jsonify([{"question": q, "count": c} for q, c in top_n])
    return jsonify([])


# D. Admin Q&A Benchmark endpoint
@app.route("/api/admin/qa_stats", methods=["GET"])
@requires_admin_auth
def get_qa_stats():
    try:
        qa_log_path = os.path.join(ADMIN_OUTPUTS, "qa_logs.json")
        if not os.path.exists(qa_log_path):
            return jsonify([])

        with open(qa_log_path, "r") as f:
            return jsonify(json.load(f))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# Celery tasks
@celery.task(bind=True)
def async_generate_media(
    self,
    media_type,
    text,
    session_id,
    api_key=None,
    template_path=None,
    aspect_ratio="16:9",
    style_type=None,
):
    """Async task for media generation"""
    try:
        # Update task state
        self.update_state(state="PROGRESS", meta={"status": "Generating..."})

        # Call the actual generation function
        output_path = generate_media(
            media_type=media_type,
            text=text,
            session_id=session_id,
            api_key=api_key,
            template_path=template_path,
            aspect_ratio=aspect_ratio,
            style_type=style_type,
        )

        return {
            "status": "SUCCESS",
            "output_path": output_path,
            "filename": os.path.basename(output_path),
        }
    except Exception as e:
        logger.error(f"Async media generation failed: {str(e)}")
        return {"status": "FAILURE", "error": str(e)}


# Modified API endpoint for async generation
@app.route("/api/generate", methods=["POST"])
def api_generate():
    data = request.form
    media_type = data.get("type")
    answer = data.get("answer", "")
    session_id = data.get("sessionId")
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")
    template_file = request.files.get("template")
    aspect_ratio = data.get("aspect_ratio", "16:9")
    style_type = data.get("style_type")

    if not api_key:
        return jsonify({"error": "Missing API key"}), 401

    try:
        # Save template if provided
        template_path = None
        if template_file and allowed_file(template_file.filename):
            filename = secure_filename(template_file.filename)
            template_path = os.path.join(UPLOAD_FOLDER, filename)
            template_file.save(template_path)

        # Start async task
        task = async_generate_media.delay(
            media_type=media_type,
            text=answer,
            session_id=session_id,
            api_key=api_key,
            template_path=template_path,
            aspect_ratio=aspect_ratio,
            style_type=style_type,
        )

        return jsonify({"task_id": task.id}), 202

    except Exception as e:
        current_app.logger.exception("Error starting generation task")
        return jsonify({"error": str(e)}), 500


# Task status endpoint
@app.route("/api/task/<task_id>", methods=["GET"])
def get_task_status(task_id):
    task = async_generate_media.AsyncResult(task_id)

    if task.state == "PENDING":
        response = {"state": task.state, "status": "Pending..."}
    elif task.state == "PROGRESS":
        response = {"state": task.state, "status": task.info.get("status", "")}
    elif task.state == "SUCCESS":
        response = {
            "state": task.state,
            "result": task.info,
            "download_url": url_for(
                "download_generated", file=task.info["filename"], _external=True
            ),
        }
    else:
        # Something went wrong
        response = {"state": task.state, "status": str(task.info)}  # Error info

    return jsonify(response)


# Admin gallery endpoints
@app.route("/api/admin/gallery", methods=["GET"])
@requires_admin_auth
def admin_gallery():
    """List generated outputs with tags"""
    outputs = []
    tags_file = os.path.join(ADMIN_OUTPUTS, "tags.json")
    tag_data = {}

    if os.path.exists(tags_file):
        try:
            with open(tags_file, "r") as f:
                tag_data = json.load(f)
        except:
            pass

    for filename in os.listdir(ADMIN_OUTPUTS):
        if filename == "tags.json":
            continue

        filepath = os.path.join(ADMIN_OUTPUTS, filename)
        if os.path.isfile(filepath):
            outputs.append(
                {
                    "filename": filename,
                    "created": datetime.fromtimestamp(
                        os.path.getctime(filepath)
                    ).isoformat(),
                    "size": os.path.getsize(filepath),
                    "tags": tag_data.get(filename, []),
                }
            )

    return jsonify(outputs)


@app.route("/api/admin/tag", methods=["POST"])
@requires_admin_auth
def admin_tag():
    """Update tags for a generated output"""
    data = request.json
    filename = data.get("filename")
    tags = data.get("tags", [])

    if not filename:
        return jsonify({"error": "Missing filename"}), 400

    filepath = os.path.join(ADMIN_OUTPUTS, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "File not found"}), 404

    # Validate tags
    valid_tags = ["usable", "inaccurate", "needs fix"]
    if any(tag not in valid_tags for tag in tags):
        return jsonify({"error": "Invalid tag"}), 400

    update_tags(filename, tags)
    return jsonify({"status": "success"})


@app.route("/api/download_generated", methods=["GET"])
def download_generated():
    file_name = request.args.get("file")
    if not file_name:
        return jsonify({"error": "Missing file"}), 400

    fp = os.path.join(OUTPUT_FOLDER, file_name)
    if not os.path.exists(fp):
        return jsonify({"error": "Not found"}), 404

    return send_file(fp, as_attachment=True)


@app.route("/api/adjust_answer", methods=["POST"])
def adjust_answer():
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not api_key:
        return jsonify({"error": "Missing API key"}), 401

    try:
        data = request.get_json()
        answer = data.get("answer")
        adj_type = data.get("type", "shorten").lower()

        if not answer or not adj_type:
            return jsonify({"error": "Missing answer or type"}), 400

        # Build prompt for the supported adjustment types
        if adj_type == "shorten":
            messages = [{"role": "user", "content": f"Shorten this text:\n\n{answer}"}]
        elif adj_type == "elaborate":
            messages = [{"role": "user", "content": f"Elaborate on this:\n\n{answer}"}]
        elif adj_type == "reword":
            messages = [{"role": "user", "content": f"Rephrase this:\n\n{answer}"}]
        else:
            return jsonify({"error": "Invalid adjustment type"}), 400

        # Call OpenAI API
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=1024,
        )
        adjusted_answer = response.choices[0].message.content.strip()

        return jsonify({"adjusted_answer": adjusted_answer})

    except Exception as e:
        logger.error(f"Error in adjust_answer: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate_from_enhanced", methods=["POST"])
def generate_from_enhanced():
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not api_key:
        return jsonify({"error": "Missing API key"}), 401

    try:
        # Enhanced parameter validation
        data = request.get_json() if request.is_json else request.form
        media_type = data.get("type")
        prompt = data.get("prompt")
        session_id = data.get("sessionId")
        aspect_ratio = data.get("aspect_ratio", "16:9")
        style_type = data.get("style_type")

        if not media_type or not prompt:
            return jsonify({"error": "Missing media type or prompt"}), 400

        if aspect_ratio not in VALID_RATIOS:
            return (
                jsonify(
                    {"error": f"Invalid aspect ratio. Must be one of: {VALID_RATIOS}"}
                ),
                400,
            )

        if style_type and style_type.upper() not in VALID_STYLES:
            return (
                jsonify(
                    {"error": f"Invalid style type. Must be one of: {VALID_STYLES}"}
                ),
                400,
            )

        # Handle template file (either base64 or file upload)
        template_path = None
        if request.is_json and data.get("template"):
            try:
                template_base64 = (
                    data["template"].split("base64,")[1]
                    if "base64," in data["template"]
                    else data["template"]
                )
                file_bytes = base64.b64decode(template_base64)
                filename = f"template_{session_id or 'temp'}.pptx"
                template_path = os.path.join(UPLOAD_FOLDER, filename)
                with open(template_path, "wb") as f:
                    f.write(file_bytes)
            except Exception as e:
                logger.error(f"Error processing base64 template: {str(e)}")
                return jsonify({"error": "Invalid template file"}), 400
        elif "template" in request.files:
            template_file = request.files["template"]
            if template_file and allowed_file(template_file.filename):
                filename = secure_filename(template_file.filename)
                template_path = os.path.join(UPLOAD_FOLDER, filename)
                template_file.save(template_path)

        # Version-pinned retry logic for Replicate models
        max_retries = 2
        retry_count = 0
        last_error = None

        while retry_count <= max_retries:
            try:
                output_path = generate_media(
                    media_type=media_type,
                    text=prompt,
                    session_id=session_id,
                    api_key=api_key,
                    template_path=template_path,
                    aspect_ratio=aspect_ratio,
                    style_type=style_type,
                )
                break
            except ReplicateError as e:
                if e.status == 422 and retry_count < max_retries:
                    retry_count += 1
                    logger.warning(f"Version mismatch, retry attempt {retry_count}")
                    continue
                last_error = str(e)
                break
            except Exception as e:
                last_error = str(e)
                break

        if not output_path or not os.path.isfile(output_path):
            logger.error(f"Generation failed: {last_error}")
            return jsonify({"error": f"Generation failed: {last_error}"}), 500

        # Save to admin outputs
        filename = os.path.basename(output_path)
        admin_output_path = os.path.join(ADMIN_OUTPUTS, filename)
        shutil.copy(output_path, admin_output_path)

        # Return download URL
        download_url = url_for("download_generated", file=filename, _external=True)

        return jsonify(
            {
                "filename": filename,
                "url": download_url,
                "metadata": {
                    "media_type": media_type,
                    "aspect_ratio": aspect_ratio,
                    "style_type": style_type,
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                },
            }
        )

    except Exception as e:
        logger.error(f"Error in generate_from_enhanced: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/enhance_prompt", methods=["POST"])
def enhance_prompt():
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not api_key:
        return jsonify({"error": "Missing API key"}), 401

    data = request.json
    prompt = data.get("prompt")
    media_type = data.get("mediaType")
    style = data.get("style")

    if not prompt:
        return jsonify({"error": "Missing prompt"}), 400

    try:
        client = OpenAI(api_key=api_key)
        enhancement_prompt = f"""
        Enhance the following prompt for {media_type} generation with a {style} style.
        The prompt should be clear, detailed, and optimized for AI generation.
        Return only the enhanced prompt, no additional commentary.
        
        Original prompt: {prompt}
        """

        # Use gpt-4-turbo model
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": enhancement_prompt}],
            temperature=0.7,
            max_tokens=500,
        )

        enhanced_prompt = response.choices[0].message.content

        # Generate preview
        preview = get_preview(media_type, enhanced_prompt, api_key)

        return jsonify({"enhanced_prompt": enhanced_prompt, "preview": preview})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


def get_preview(media_type, prompt, api_key):
    client = OpenAI(api_key=api_key)

    if media_type == "image":
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "user", "content": f"Describe a thumbnail image for: {prompt}"}
            ],
            max_tokens=150,
        )
        return response.choices[0].message.content

    elif media_type == "slides":
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {
                    "role": "user",
                    "content": f"Create an outline for slides based on: {prompt}",
                }
            ],
            max_tokens=300,
        )
        return response.choices[0].message.content

    elif media_type == "video":
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {
                    "role": "user",
                    "content": f"Describe a short video preview for: {prompt}",
                }
            ],
            max_tokens=200,
        )
        return response.choices[0].message.content

    return "Preview not available for this media type"


# Main route
@app.route("/")
def index():
    roles = [
        {"value": "general", "label": "General"},
        {"value": "engineering", "label": "Engineering"},
        {"value": "marketing", "label": "Marketing"},
        {"value": "public", "label": "Public"},
    ]
    return render_template("kaggle.html", roles=roles)


@app.route("/api/debug_vector", methods=["POST"])
def debug_vector():
    """Debug endpoint for testing vector embeddings"""
    data = request.json
    text = data.get("text")
    model = data.get("model", "text-embedding-3-large")
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")

    if not text or not api_key:
        return jsonify({"error": "Missing text or API key"}), 400

    try:
        embedding = embed(text, api_key, model)
        return jsonify(
            {
                "text": text,
                "embedding": embedding[:5] + ["..."],  # Show partial for debugging
                "dimensions": len(embedding),
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate_preview", methods=["POST"])
def generate_preview():
    """Generate preview for enhanced prompts with improved media generation"""
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not api_key:
        return jsonify({"error": "Missing API key"}), 401

    try:
        data = request.json
        prompt = data.get("prompt")
        media_type = data.get("mediaType")
        aspect_ratio = data.get("aspect_ratio", "16:9")
        style_type = data.get("style_type")

        if not prompt or not media_type:
            return jsonify({"error": "Missing prompt or media type"}), 400

        # Enhanced prompt engineering
        client = OpenAI(api_key=api_key)

        if media_type == "image":
            # Improved image preview with style considerations
            preview_prompt = f"""
            Generate a detailed thumbnail description for an image with these characteristics:
            - Style: {style_type if style_type else 'general'}
            - Aspect ratio: {aspect_ratio}
            - Based on this content: {prompt}
            
            The description should be 1-2 sentences focusing on key visual elements.
            """
            response = client.chat.completions.create(
                model="gpt-4-turbo",
                messages=[{"role": "user", "content": preview_prompt}],
                max_tokens=150,
            )
            return jsonify({"preview": response.choices[0].message.content})

        elif media_type == "slides":
            # Enhanced slides outline with structure
            outline_prompt = f"""
            Create a detailed slides outline for a presentation with:
            - Professional structure
            - Balanced content distribution
            - Visual elements suggestions
            - Based on this content: {prompt}
            
            Return as a markdown list with slide titles and bullet points.
            """
            response = client.chat.completions.create(
                model="gpt-4-turbo",
                messages=[{"role": "user", "content": outline_prompt}],
                max_tokens=300,
            )
            return jsonify({"outline": response.choices[0].message.content})

        elif media_type == "video":
            # Improved video preview with scene breakdown
            video_prompt = f"""
            Describe a short video preview with:
            - Key scenes (3-5)
            - Suggested visuals
            - Pacing notes
            - Based on this content: {prompt}
            
            Keep it under 200 words.
            """
            response = client.chat.completions.create(
                model="gpt-4-turbo",
                messages=[{"role": "user", "content": video_prompt}],
                max_tokens=250,
            )
            return jsonify({"preview": response.choices[0].message.content})

        return jsonify({"preview": "Preview not available for this media type"})

    except Exception as e:
        logger.error(f"Error in generate_preview: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate_slides_outline", methods=["POST"])
def generate_slides_outline():
    """Generate slides outline with improved JSON structure"""
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not api_key:
        return jsonify({"error": "Missing API key"}), 401

    try:
        data = request.json
        prompt = data.get("prompt")
        depth = data.get("depth", "detailed")  # 'brief' or 'detailed'

        if not prompt:
            return jsonify({"error": "Missing prompt"}), 400

        client = OpenAI(api_key=api_key)

        # Improved structured prompt
        outline_prompt = f"""
        Create a {depth} slides outline in JSON format based on: {prompt}
        
        Return a JSON object with this structure:
        {{
            "title": "Presentation Title",
            "slides": [
                {{
                    "title": "Slide Title",
                    "content": ["bullet point 1", "bullet point 2"],
                    "visual_suggestion": "suggested visual element",
                    "notes": "additional speaker notes"
                }}
            ],
            "recommendations": {{
                "design": "suggested design style",
                "timing": "estimated presentation duration",
                "audience": "target audience considerations"
            }}
        }}
        """

        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": outline_prompt}],
            response_format={"type": "json_object"},
            max_tokens=1000,
        )

        outline = json.loads(response.choices[0].message.content)

        # Validate and clean the response
        if not isinstance(outline, dict):
            raise ValueError("Invalid outline format")

        if "slides" not in outline:
            outline["slides"] = []

        return jsonify(outline)

    except Exception as e:
        logger.error(f"Error in generate_slides_outline: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/feedback", methods=["POST"])
def handle_feedback():
    data = request.json
    feedback = {
        "rating": data.get("rating"),
        "comment": data.get("comment", ""),
        "media_type": data.get("mediaType", ""),
        "prompt": data.get("prompt", ""),
        "style": data.get("style", ""),
        "reference_image": bool(data.get("hasReference", False)),
        "session_id": data.get("sessionId"),
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }

    # Save feedback to file
    feedback_path = os.path.join(ADMIN_OUTPUTS, "feedback.json")
    feedbacks = []

    if os.path.exists(feedback_path):
        with open(feedback_path, "r") as f:
            try:
                feedbacks = json.load(f)
            except:
                pass

    feedbacks.append(feedback)

    with open(feedback_path, "w") as f:
        json.dump(feedbacks, f, indent=2)

    return jsonify({"status": "success"})


@app.route("/api/admin/feedback", methods=["GET"])
@requires_admin_auth
def get_feedback():
    feedback_path = os.path.join(ADMIN_OUTPUTS, "feedback.json")
    if os.path.exists(feedback_path):
        with open(feedback_path, "r") as f:
            return jsonify(json.load(f))
    return jsonify([])


@app.route("/api/admin/media_summary", methods=["GET"])
@requires_admin_auth
def media_summary():
    feedback_path = os.path.join(ADMIN_OUTPUTS, "feedback.json")
    if not os.path.exists(feedback_path):
        return jsonify({"summary": "No feedback data available"})

    with open(feedback_path, "r") as f:
        feedbacks = json.load(f)

    # Generate summary using GPT
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        prompt = f"""
        Analyze this feedback data and provide a concise summary of common suggestions:
        {json.dumps(feedbacks[:20])}  # Limit to first 20 for context
        """

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
        )
        return jsonify({"summary": response.choices[0].message.content})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/admin/export_feedback", methods=["GET"])
@requires_admin_auth
def export_feedback():
    fmt = request.args.get("format", "txt")
    feedback_path = os.path.join(ADMIN_OUTPUTS, "feedback.json")

    if not os.path.exists(feedback_path):
        return jsonify({"error": "No feedback data"}), 404

    with open(feedback_path, "r") as f:
        feedbacks = json.load(f)

    if fmt == "txt":
        content = "\n".join(
            f"[{f['timestamp']}] {f['rating']} stars: {f['comment']}" for f in feedbacks
        )
        return content, 200, {"Content-Type": "text/plain"}

    elif fmt == "docx":
        doc = DocxWriter()
        for fb in feedbacks:
            doc.add_paragraph(f"Rating: {fb['rating']} stars")
            doc.add_paragraph(f"Comment: {fb['comment']}")
            doc.add_paragraph(f"Media Type: {fb.get('media_type', 'N/A')}")
            doc.add_paragraph("")
        doc_path = "feedback.docx"
        doc.save(doc_path)
        return send_file(doc_path, as_attachment=True)

    elif fmt == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for fb in feedbacks:
            pdf.cell(200, 10, txt=f"Rating: {fb['rating']} stars", ln=1)
            pdf.multi_cell(0, 10, txt=f"Comment: {fb['comment']}")
            pdf.ln(5)
        pdf_path = "feedback.pdf"
        pdf.output(pdf_path)
        return send_file(pdf_path, as_attachment=True)

    return jsonify({"error": "Invalid format"}), 400


def fetch_competitor_analysis(api_key):
    client = OpenAI(api_key=api_key)

    competitors_prompt = """
You are a competitive intelligence analyst specializing in MPC (Multi-Party Computation) technology. 
Analyze Crossbar Cramium Lab's MPC SDK and identify at least 4 real competitors with comparable products.

For each competitor, provide:
1. Their core MPC offering
2. Supported blockchains (count)
3. Key differentiating features
4. Security certifications
5. Pricing model (if available)
6. Notable customers/partners

Crossbar Cramium Lab's key features to compare against:
- Threshold signature schemes
- Secure key generation
- Multi-cloud support
- Hardware security module integration
- Compliance certifications
- Developer tools and SDKs

Return only a JSON list like:
{
  "competitors": [
    {
      "name": "Fireblocks",
      "blockchains": 24,
      "features": ["Institutional-grade security", "Policy engine", "DeFi connectivity"],
      "certifications": ["SOC 2 Type II", "ISO 27001"],
      "pricing": "Volume-based",
      "customers": ["Binance", "Crypto.com"],
      "comparison": "Stronger in institutional features but less flexible in key management"
    },
    ...
  ]
}
Only include verifiable information from public sources or reasonable estimates.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": competitors_prompt}],
            response_format={"type": "json_object"},
            temperature=0.2,
            max_tokens=1000,
        )
        content = response.choices[0].message.content
        response_data = json.loads(content)
        return response_data.get("competitors", [])
    except Exception as e:
        logger.error(f"Failed to fetch competitor analysis: {e}")
        return []


@app.route("/api/visualize", methods=["POST"])
def visualize_data():
    data = request.json
    text = data.get("text")
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")

    if not text or not api_key:
        return jsonify({"error": "Missing text or API key"}), 400

    try:
        # Get comprehensive competitor analysis
        competitors = fetch_competitor_analysis(api_key)
        if competitors:
            text += "\n\nCompetitor Analysis:\n"
            for c in competitors:
                text += f"\n{c.get('name', 'Unnamed competitor')}:\n"
                text += f"- Blockchains supported: {c.get('blockchains', 'N/A')}\n"
                text += f"- Key features: {', '.join(c.get('features', []))}\n"
                text += f"- Certifications: {', '.join(c.get('certifications', []))}\n"
                text += f"- Pricing: {c.get('pricing', 'Unknown')}\n"
                text += f"- Notable customers: {', '.join(c.get('customers', []))}\n"
                text += (
                    f"- Comparison: {c.get('comparison', 'No comparison available')}\n"
                )

        # Updated visualization prompt with specific axis instructions
        client = OpenAI(api_key=api_key)
        prompt = f"""
Analyze this competitor comparison data and create a visualization specification.
The chart should clearly show comparisons between different competitors.

Key requirements:
1. Competitor names must appear on the x-axis
2. Use bar charts as the primary visualization
3. Each bar should represent a different competitor
4. Include Crossbar Cramium Lab in the comparison if data is available
5. Possible metrics to visualize:
   - Number of supported blockchains
   - Count of security certifications
   - Number of key features
   - Count of notable customers

Structure the Vega-Lite spec as follows:
- x-axis: "name" (competitor names)
- y-axis: choose the most meaningful quantitative metric
- color: use a distinct color for each competitor
- title: "Competitive Analysis: [metric being compared]"

Example structure:
{{
  "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
  "description": "Competitor comparison",
  "data": {{"values": [/* competitor data */]}},
  "mark": "bar",
  "encoding": {{
    "x": {{"field": "name", "type": "nominal", "title": "Competitor"}},
    "y": {{"field": "[metric]", "type": "quantitative", "title": "[Metric Name]"}},
    "color": {{"field": "name", "legend": null}}
  }},
  "title": "Competitive Analysis: [metric]"
}}

Text to analyze:
{text}
"""

        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            max_tokens=1500,
        )

        content = response.choices[0].message.content
        spec = json.loads(content)
        return jsonify(spec)

    except Exception as e:
        logger.error(f"Error in visualize_data: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/search", methods=["POST"])
def internet_search():
    data = request.json
    query = data.get("query")
    api_key = request.headers.get("Authorization", "").replace("Bearer ", "")

    if not query or not api_key:
        return jsonify({"error": "Missing query or API key"}), 400

    try:
        client = OpenAI(api_key=api_key)

        # Use GPT to generate synthetic search results
        prompt = f"""
        You are a web search assistant. Given the query "{query}", generate 3-5 realistic 
        search results with titles, URLs, and snippets as if they came from a real search engine.
        
        Return the results in JSON format like this:
        {{
            "results": [
                {{
                    "title": "Result title",
                    "link": "https://example.com/relevant-page",
                    "snippet": "Relevant information snippet..."
                }}
            ]
        }}
        """

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            max_tokens=1000,
        )

        search_results = json.loads(response.choices[0].message.content)

        return jsonify(
            {"status": "success", "results": search_results.get("results", [])}
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    print("🔧 Starting Crossbar Flask app on http://localhost:5000")
    try:
        load_context_index()
        app.run(host="0.0.0.0", port=5000, debug=True)
    except Exception as e:
        print("❌ Failed to start the Flask app:", e)
